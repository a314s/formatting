import os
import re
import subprocess
import pandas as pd
import cv2
import numpy as np
from PIL import Image
import imagehash
import io
import string
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
try:
    from docx.enum.table import WD_TABLE_LAYOUT
except ImportError:
    WD_TABLE_LAYOUT = None
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from google.cloud import speech_v1p1beta1 as speech
from google.cloud.speech_v1p1beta1.types import WordInfo
from google.oauth2 import service_account
from google.cloud import storage
import math

# Helper functions
def add_field_code(run, field_code):
    """Add a Word field code to the given run"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(begin)
    
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = field_code
    run._r.append(instr)
    
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.append(end)

def extract_audio(video_file, audio_file):
    """Extract audio from video file using ffmpeg"""
    command = [
        'ffmpeg',
        '-y',
        '-i', video_file,
        '-vn',
        '-acodec', 'pcm_s16le',
        '-f', 'wav',
        audio_file
    ]
    result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if result.returncode != 0:
        raise Exception(f"FFmpeg Error: {result.stderr}")
    return audio_file

def get_audio_properties(audio_file):
    """Get audio properties using ffprobe"""
    command = ['ffprobe', '-v', 'error', '-show_entries', 'stream=sample_rate,channels', 
               '-of', 'default=noprint_wrappers=1', audio_file]
    result = subprocess.run(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    output = result.stdout
    properties = {}
    for line in output.strip().split('\n'):
        if 'sample_rate' in line:
            properties['sample_rate'] = int(line.split('=')[1])
        elif 'channels' in line:
            properties['channels'] = int(line.split('=')[1])
    return properties

def upload_to_gcs(audio_file, bucket_name, credential_file):
    """Upload audio file to Google Cloud Storage"""
    credentials = service_account.Credentials.from_service_account_file(credential_file)
    client = storage.Client(credentials=credentials)

    bucket = client.bucket(bucket_name)
    blob_name = os.path.basename(audio_file)
    blob = bucket.blob(blob_name)
    blob.upload_from_filename(audio_file)

    gcs_uri = f"gs://{bucket_name}/{blob_name}"
    return gcs_uri

def transcribe_audio_gcs(gcs_uri, credential_file, sample_rate_hertz, audio_channel_count):
    """Transcribe audio using Google Cloud Speech-to-Text"""
    credentials = service_account.Credentials.from_service_account_file(credential_file)
    client = speech.SpeechClient(credentials=credentials)

    audio = speech.RecognitionAudio(uri=gcs_uri)

    config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=sample_rate_hertz,
        language_code='en-US',
        enable_word_time_offsets=True,
        enable_automatic_punctuation=True,
        audio_channel_count=audio_channel_count,
        enable_separate_recognition_per_channel=False,
    )

    operation = client.long_running_recognize(config=config, audio=audio)
    response = operation.result(timeout=600)

    transcription_words = []

    for result in response.results:
        alternative = result.alternatives[0]
        for word_info in alternative.words:
            transcription_words.append(word_info)

    return transcription_words

def load_script(excel_file):
    """Load script from Excel file"""
    # Read the Excel file without headers
    df = pd.read_excel(excel_file, header=None)
    # Assuming the script lines are in the first column (column 0)
    df.columns = ['Script']
    script_lines = df['Script'].tolist()
    return script_lines, df

def align_script_with_transcription(script_lines, transcription_words):
    """Align script with transcription using dynamic programming"""
    import numpy as np
    import re

    # Preprocess transcription words
    transcription_word_list = []
    for tw in transcription_words:
        word_text = tw.word.lower()
        word_clean = re.sub(r'[^\w\s]', '', word_text)
        transcription_word_list.append({
            'word': word_clean,
            'start_time': tw.start_time.total_seconds(),
            'end_time': tw.end_time.total_seconds(),
        })

    # Preprocess script words with line numbers
    script_word_list = []
    line_indices = []
    for idx, line in enumerate(script_lines):
        line_clean = re.sub(r'[^\w\s]', '', line.lower())
        words = line_clean.strip().split()
        for word in words:
            script_word_list.append({'word': word, 'line_number': idx})
            line_indices.append(idx)

    # Initialize scoring matrices
    n = len(script_word_list)
    m = len(transcription_word_list)
    score_matrix = np.zeros((n + 1, m + 1))
    traceback_matrix = np.zeros((n + 1, m + 1), dtype=int)

    # Scoring parameters
    match_score = 1
    mismatch_penalty = -1
    gap_penalty = -1

    # Fill the score and traceback matrices
    for i in range(1, n + 1):
        score_matrix[i, 0] = score_matrix[i - 1, 0] + gap_penalty
        traceback_matrix[i, 0] = 1  # Up
    for j in range(1, m + 1):
        score_matrix[0, j] = score_matrix[0, j - 1] + gap_penalty
        traceback_matrix[0, j] = 2  # Left

    for i in range(1, n + 1):
        script_word = script_word_list[i - 1]['word']
        for j in range(1, m + 1):
            transcription_word = transcription_word_list[j - 1]['word']
            if script_word == transcription_word:
                diag_score = score_matrix[i - 1, j - 1] + match_score
            else:
                diag_score = score_matrix[i - 1, j - 1] + mismatch_penalty
            up_score = score_matrix[i - 1, j] + gap_penalty
            left_score = score_matrix[i, j - 1] + gap_penalty

            max_score = max(diag_score, up_score, left_score)
            score_matrix[i, j] = max_score

            if max_score == diag_score:
                traceback_matrix[i, j] = 3  # Diagonal
            elif max_score == up_score:
                traceback_matrix[i, j] = 1  # Up
            else:
                traceback_matrix[i, j] = 2  # Left

    # Traceback to find the alignment
    i = n
    j = m
    alignment = []
    while i > 0 or j > 0:
        if traceback_matrix[i, j] == 3:
            alignment.append((i - 1, j - 1))  # Match/Mismatch
            i -= 1
            j -= 1
        elif traceback_matrix[i, j] == 1:
            alignment.append((i - 1, None))  # Deletion (gap in transcription)
            i -= 1
        else:
            alignment.append((None, j - 1))  # Insertion (gap in script)
            j -= 1
    alignment.reverse()

    # Collect timestamps for each script line
    line_times = {}
    for (script_idx, trans_idx) in alignment:
        if script_idx is not None and trans_idx is not None:
            line_number = script_word_list[script_idx]['line_number']
            start_time = transcription_word_list[trans_idx]['start_time']
            end_time = transcription_word_list[trans_idx]['end_time']
            if line_number not in line_times:
                line_times[line_number] = {'start_time': start_time, 'end_time': end_time}
            else:
                line_times[line_number]['start_time'] = min(start_time, line_times[line_number]['start_time'])
                line_times[line_number]['end_time'] = max(end_time, line_times[line_number]['end_time'])

    # Build script_alignment
    script_alignment = []
    for idx, line in enumerate(script_lines):
        if idx in line_times:
            start_time = line_times[idx]['start_time']
            end_time = line_times[idx]['end_time']
        else:
            start_time = None
            end_time = None
        script_alignment.append({'line': line, 'start_time': start_time, 'end_time': end_time})

    return script_alignment

def extract_frames(video_file, frames_folder):
    """Extract frames from video using color detection and image hashing"""
    if not os.path.exists(frames_folder):
        os.makedirs(frames_folder)

    cap = cv2.VideoCapture(video_file)
    extracted_frames = 0
    frame_data = []  # List to store frame timestamps and filenames
    last_hash = None

    # Define the target color in BGR
    target_color_bgr = np.array([52, 116, 255])  # #ff7434 in BGR
    # Allowable color range (adjust as needed)
    lower_bound = np.array([40, 100, 240])
    upper_bound = np.array([60, 130, 255])

    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break

        # Check if the frame is all black
        if not np.any(frame):
            # Frame is all black, skip it
            continue

        # Create a mask for the target color
        mask = cv2.inRange(frame, lower_bound, upper_bound)
        if cv2.countNonZero(mask) > 0:
            # Compute hash of the frame
            pil_image = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
            current_hash = imagehash.average_hash(pil_image)
            # Compare with last hash
            is_duplicate = False
            if last_hash is not None:
                difference = current_hash - last_hash
                if difference < 5:  # Threshold for similarity (adjust as needed)
                    is_duplicate = True

            if not is_duplicate:
                # Save the frame
                timestamp = cap.get(cv2.CAP_PROP_POS_MSEC) / 1000.0  # Convert to seconds
                minutes = int(timestamp // 60)
                seconds = int(timestamp % 60)
                tenths = int((timestamp - minutes * 60 - seconds) * 10)
                frame_filename = os.path.join(frames_folder, f'frame_{minutes:02d}_{seconds:02d}_{tenths}.png')
                cv2.imwrite(frame_filename, frame)
                # Save the timestamp and filename and hash
                frame_data.append({'timestamp': timestamp, 'filename': frame_filename, 'hash': current_hash})
                extracted_frames += 1
                # Update last hash
                last_hash = current_hash

            # Skip ahead to avoid extracting multiple frames for the same scene
            cap.set(cv2.CAP_PROP_POS_FRAMES, cap.get(cv2.CAP_PROP_POS_FRAMES) + 30)

    cap.release()
    return frame_data

def extract_additional_frames(video_file, script_alignment, existing_frame_data, frames_folder):
    """Extract additional frames based on the end times of script lines"""
    cap = cv2.VideoCapture(video_file)
    additional_frames = 0

    fps = cap.get(cv2.CAP_PROP_FPS)
    total_frames = cap.get(cv2.CAP_PROP_FRAME_COUNT)
    if not fps or not total_frames:
        cap.release()
        return existing_frame_data

    video_duration = total_frames / fps

    existing_timestamps = [f['timestamp'] for f in existing_frame_data]
    existing_hashes = [f['hash'] for f in existing_frame_data if 'hash' in f]
    existing_timestamps.sort()

    # Build desired_timestamps using end_time + 2.0 seconds
    desired_timestamps = [1.0]  # still include the initial 1 second frame if needed
    for item in script_alignment:
        end_time = item['end_time']
        if end_time is not None:
            adjusted_time = end_time + 2.0
            if adjusted_time > video_duration:
                adjusted_time = video_duration - 0.1
            desired_timestamps.append(adjusted_time)

    desired_timestamps = sorted(set(desired_timestamps))

    for timestamp in desired_timestamps:
        # Check if close to existing frames within 1 second
        index = bisect_left(existing_timestamps, timestamp)
        close = False
        if index < len(existing_timestamps) and abs(existing_timestamps[index] - timestamp) <= 1.0:
            close = True
        if index > 0 and not close and abs(existing_timestamps[index - 1] - timestamp) <= 1.0:
            close = True
        if close:
            continue

        # Extract frame at this timestamp
        cap.set(cv2.CAP_PROP_POS_MSEC, timestamp * 1000)
        ret, frame = cap.read()
        if not ret:
            continue

        # Check if frame is all black
        if not np.any(frame):
            continue

        pil_image = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
        current_hash = imagehash.average_hash(pil_image)

        # Duplicate check with 3-second rule
        is_duplicate = False
        for h_idx, e_hash in enumerate(existing_hashes):
            if (current_hash - e_hash) == 0:
                existing_time = existing_timestamps[h_idx]
                if abs(timestamp - existing_time) <= 3.0:
                    is_duplicate = True
                    break

        if is_duplicate:
            continue
        else:
            minutes = int(timestamp // 60)
            seconds = int(timestamp % 60)
            tenths = int((timestamp - minutes * 60 - seconds) * 10)
            frame_filename = os.path.join(frames_folder, f'frame_{minutes:02d}_{seconds:02d}_{tenths}.png')
            cv2.imwrite(frame_filename, frame)
            existing_frame_data.append({'timestamp': timestamp, 'filename': frame_filename, 'hash': current_hash})
            existing_timestamps.append(timestamp)
            existing_hashes.append(current_hash)
            existing_timestamps.sort()
            additional_frames += 1

    cap.release()
    return existing_frame_data

# Import bisect for extract_additional_frames
from bisect import bisect_left

def assign_steps_to_frames(frame_data, script_alignment):
    """Assign steps to frames based on end_time"""
    # Sort frames by their timestamp
    frame_data.sort(key=lambda x: x['timestamp'])

    frame_step_mapping = []
    num_frames = len(frame_data)
    num_scripts = len(script_alignment)
    script_index = 0

    # Sort script lines by end_time to efficiently map them
    script_alignment_sorted = sorted(script_alignment, key=lambda x: (float('inf') if x['end_time'] is None else x['end_time']))

    # We'll iterate through frames and assign lines based on end_time intervals
    line_idx = 0
    for i in range(num_frames):
        frame_time = frame_data[i]['timestamp']
        next_frame_time = frame_data[i+1]['timestamp'] if i+1 < num_frames else math.inf

        frame_lines = []

        # Advance line_idx until we find lines with end_time >= frame_time
        while line_idx < num_scripts:
            end_time = script_alignment_sorted[line_idx]['end_time']
            # If no end_time, skip line
            if end_time is None:
                line_idx += 1
                continue

            # If this line ends after the end of this frame interval, no more lines for this frame
            if end_time >= next_frame_time:
                break

            if end_time >= frame_time and end_time < next_frame_time:
                # This line belongs to this frame
                frame_lines.append(script_alignment_sorted[line_idx]['line'])
                line_idx += 1
            elif end_time < frame_time:
                # This line ended too early, move on
                line_idx += 1
            else:
                # end_time >= next_frame_time means we've passed this frame
                break

        # Combine text lines into a single text block
        combined_text = '\n'.join(frame_lines)

        frame_step_mapping.append({
            'frame': frame_data[i]['filename'],
            'text': combined_text
        })

    return frame_step_mapping

def set_cell_border(cell, **kwargs):
    """Set cell's border"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    tcPr.append(tcBorders)
    
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_element = OxmlElement(f'w:{edge}')
            tcBorders.append(edge_element)
            for key in ['sz', 'val', 'color', 'space', 'shadow']:
                if key in edge_data:
                    edge_element.set(qn(f'w:{key}'), str(edge_data[key]))

def add_spacing_table(document, height_cm):
    """Add invisible table to create precise vertical spacing"""
    from docx.shared import Cm
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    
    # Create a single-cell table for spacing
    table = document.add_table(rows=1, cols=1)
    
    # Remove borders
    for cell in table.rows[0].cells:
        set_cell_border(
            cell,
            top={"sz": "0", "val": "none"},
            bottom={"sz": "0", "val": "none"},
            start={"sz": "0", "val": "none"},
            end={"sz": "0", "val": "none"}
        )
    
    # Set fixed height
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    table.rows[0].height = Cm(height_cm)
    
    # Remove cell padding and spacing
    cell = table.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    
    # Remove any existing margins
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>' +
                      f'<w:top w:w="0" w:type="dxa"/>' +
                      f'<w:bottom w:w="0" w:type="dxa"/>' +
                      f'<w:start w:w="0" w:type="dxa"/>' +
                      f'<w:end w:w="0" w:type="dxa"/>' +
                      f'</w:tcMar>')
    tcPr.append(tcMar)
    
    return table

def create_header_tables(section, video_file, assets_folder):
    """Create header using two-row table with merged cells"""
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls
    from datetime import datetime
    import os
    
    # Define a mid gray color to use consistently
    mid_gray = RGBColor(102, 102, 102)  # A nice medium gray (hex: #666666)
    
    # Use full page width
    total_width_cm = section.page_width.cm
    header_width_cm = total_width_cm
    
    # Handle first page header
    if section.first_page_header:
        header = section.first_page_header
        # Clear existing content
        for element in list(header._element):
            header._element.remove(element)
        
        # Create first page header table
        table = header.add_table(rows=2, cols=3, width=Cm(header_width_cm))
        table.allow_autofit = False
        table.style = 'Table Grid'
        
        # Set column widths
        widths = [0.33, 0.34, 0.33]
        for i, width in enumerate(widths):
            table.columns[i].width = Cm(total_width_cm * width)
        
        # Left column: Logo
        logo_cell = table.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = logo_paragraph.add_run()
        try:
            # Calculate cell width and height for logo with increased height
            cell_width = table.columns[0].width - Pt(4)
            cell_height = Pt(70) - Pt(4)  # Adjusted to match new row height
            
            logo_path = os.path.join(assets_folder, "logo.jpg")
            from PIL import Image
            with Image.open(logo_path) as img:
                aspect_ratio = img.width / img.height
            
            target_height = min(Pt(60), cell_height)  # Increased from 30 to 60
            target_width = min(target_height * aspect_ratio, cell_width)
            
            final_width = min(target_width, cell_width)
            final_height = min(target_height, final_width / aspect_ratio)
            
            logo_run.add_picture(logo_path, width=final_width, height=final_height)
        except Exception as e:
            print(f"Warning: Could not add logo: {str(e)}")
        
        # Center column: Split into two parts vertically (Vendor ID and Document ID)
        center_cell = table.cell(0, 1)
        
        # First part: Vendor ID (top)
        center_paragraph1 = center_cell.paragraphs[0]
        center_paragraph1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        center_paragraph1.space_after = Pt(12) # Add space between paragraphs
        vendor_run = center_paragraph1.add_run('Vendor ID: 3022360')
        vendor_run.font.name = 'Arial'
        vendor_run.font.size = Pt(12)
        vendor_run.font.bold = False
        vendor_run.font.color.rgb = mid_gray
        
        # Second part: Document ID (bottom)
        center_paragraph2 = center_cell.add_paragraph()
        center_paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc_id_run = center_paragraph2.add_run('Document ID: TDxxxxxx')
        doc_id_run.font.name = 'Arial'
        doc_id_run.font.size = Pt(12)
        doc_id_run.font.bold = False
        doc_id_run.font.color.rgb = mid_gray
        
        # Right column: Page number and address
        right_cell = table.cell(0, 2)
        # Change vertical alignment to TOP for the address column
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        right_paragraph = right_cell.paragraphs[0]
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Changed to RIGHT align
        
        # Add page number
        page_run = right_paragraph.add_run('Page ')
        page_run.font.name = 'Arial'
        page_run.font.size = Pt(12)
        page_run.font.bold = False
        page_run.font.color.rgb = mid_gray
        add_field_code(page_run, 'PAGE')
        page_run = right_paragraph.add_run(' of ')
        page_run.font.name = 'Arial'
        page_run.font.size = Pt(12)
        page_run.font.bold = False
        page_run.font.color.rgb = mid_gray
        add_field_code(page_run, 'NUMPAGES')
        
        # Add a new paragraph for address with left alignment
        addr_paragraph = right_cell.add_paragraph()
        addr_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add address on new lines
        address_lines = [
            "\nRamat Gavriel Industrial Park",
            "P. O. Box 544",
            "Migdal Ha'Emek 23150, Israel",
            "Telephone: 972-(0)4-604-8100",
            "http://www.Camtek.com"
        ]
        
        for line in address_lines:
            addr_run = addr_paragraph.add_run(line + '\n')
            addr_run.font.name = 'Arial'
            addr_run.font.size = Pt(10)
            addr_run.font.color.rgb = mid_gray
        
        # Set row height and borders
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[0].height = Pt(90)  # Increased from 70pt to 100pt to accommodate address text
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[1].height = Pt(30)  # Set appropriate height for the Date/Name/Signature row
        
        # Add borders
        for row in table.rows:
            for cell in row.cells:
                # Apply center alignment to all cells except the right cell in the first row
                if not (cell == right_cell):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                set_cell_border(
                    cell,
                    top={"sz": "2", "val": "single", "color": "000000"},
                    bottom={"sz": "2", "val": "single", "color": "000000"},
                    start={"sz": "2", "val": "single", "color": "000000"},
                    end={"sz": "2", "val": "single", "color": "000000"}
                )
        
        # Set up the second row with Date, Name, and Signature fields
        # Date column
        from datetime import datetime
        current_date = datetime.now().strftime("%B %d, %Y")  # Format like "February 25, 2025"
        
        date_cell = table.cell(1, 0)
        date_paragraph = date_cell.paragraphs[0]
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_run = date_paragraph.add_run('Revision Date: ' + current_date)
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(12)
        date_run.font.bold = False
        date_run.font.color.rgb = mid_gray
        
        # Name column
        name_cell = table.cell(1, 1)
        name_paragraph = name_cell.paragraphs[0]
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_paragraph.add_run('Name: ')
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(12)
        name_run.font.bold = False
        name_run.font.color.rgb = mid_gray
        
        # Signature column
        signature_cell = table.cell(1, 2)
        signature_paragraph = signature_cell.paragraphs[0]
        signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        signature_run = signature_paragraph.add_run('Signature: ')
        signature_run.font.name = 'Arial'
        signature_run.font.size = Pt(12)
        signature_run.font.bold = False
        signature_run.font.color.rgb = mid_gray
    
    # Handle regular header
    if section.header:
        header = section.header
        # Clear existing content
        for element in list(header._element):
            header._element.remove(element)
        
        # Create regular header table
        table = header.add_table(rows=2, cols=4, width=Cm(header_width_cm))
        table.allow_autofit = False
        table.style = 'Table Grid'
        
        # Set column widths (as proportions of total width)
        widths = [0.25, 0.25, 0.25, 0.25]  # Equal widths for all columns
        for i, width in enumerate(widths):
            table.columns[i].width = Cm(total_width_cm * width)
        
        # Merge first column vertically (A1 and A2)
        a1 = table.cell(0, 0)
        a2 = table.cell(1, 0)
        a1.merge(a2)
        
        # Merge remaining cells in second row (B2:D2)
        b2 = table.cell(1, 1)
        d2 = table.cell(1, 3)
        b2.merge(d2)
        
        # Add logo to A1:A2 (merged)
        logo_cell = table.cell(0, 0)
        logo_paragraph = logo_cell.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = logo_paragraph.add_run()
        try:
            # Calculate cell width and height for logo
            cell_width = table.columns[0].width - Pt(4)
            cell_height = Pt(25) * 2 - Pt(4)  # Two rows height
            
            logo_path = os.path.join(assets_folder, "logo.jpg")
            from PIL import Image
            with Image.open(logo_path) as img:
                aspect_ratio = img.width / img.height
            
            target_height = min(Pt(45), cell_height)
            target_width = min(target_height * aspect_ratio, cell_width)
            
            final_width = min(target_width, cell_width)
            final_height = min(target_height, final_width / aspect_ratio)
            
            logo_run.add_picture(logo_path, width=final_width, height=final_height)
        except Exception as e:
            print(f"Warning: Could not add logo: {str(e)}")
        
        # Add document number to B1
        title_cell = table.cell(0, 1)
        title_paragraph = title_cell.paragraphs[0]
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run('TDxxxxxxxx')
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(12)
        title_run.font.bold = False
        title_run.font.color.rgb = mid_gray
        
        # Add date to C1
        date_cell = table.cell(0, 2)
        date_paragraph = date_cell.paragraphs[0]
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_paragraph.add_run(datetime.now().strftime("%B %d, %Y"))
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(12)
        date_run.font.bold = False
        date_run.font.color.rgb = mid_gray
        
        # Add page number to D1
        number_cell = table.cell(0, 3)
        number_paragraph = number_cell.paragraphs[0]
        number_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        page_run = number_paragraph.add_run('Page ')
        page_run.font.name = 'Arial'
        page_run.font.size = Pt(12)
        page_run.font.bold = False
        page_run.font.color.rgb = mid_gray
        add_field_code(page_run, 'PAGE')
        page_run = number_paragraph.add_run(' of ')
        page_run.font.name = 'Arial'
        page_run.font.size = Pt(12)
        page_run.font.bold = False
        page_run.font.color.rgb = mid_gray
        add_field_code(page_run, 'NUMPAGES')
        
        # Add document name to merged B2:D2
        basename = os.path.basename(video_file)
        name_without_ext, _ = os.path.splitext(basename)
        doc_name = name_without_ext + '-PDF'
        
        name_cell = table.cell(1, 1)
        name_paragraph = name_cell.paragraphs[0]
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(doc_name)
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(12)
        name_run.font.bold = False
        name_run.font.color.rgb = mid_gray
        
        # Set row heights for regular header
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[0].height = Pt(25)
        table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[1].height = Pt(25)
        
        # Add borders
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(
                    cell,
                    top={"sz": "2", "val": "single", "color": "000000"},
                    bottom={"sz": "2", "val": "single", "color": "000000"},
                    start={"sz": "2", "val": "single", "color": "000000"},
                    end={"sz": "2", "val": "single", "color": "000000"}
                )

def generate_word_document(frame_step_mapping, video_file, doc_name, output_path, assets_folder):
    """Generate Word document with frames and text"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.shared import Inches, Twips
    
    document = Document()
    
    # Configure document settings
    section = document.sections[-1]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    # Reduce top margin to compensate for large header
    section.top_margin = Cm(1.5)  # Reduced from 2cm
    section.bottom_margin = Cm(1)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)
    section.different_first_page_header_footer = True
    
    # Create headers
    create_header_tables(section, video_file, assets_folder)
    
    # Compute footer width and add footers
    footer_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm

    # Add footers
    for footer_type in [section.first_page_footer, section.footer]:
        for element in list(footer_type._element):
            footer_type._element.remove(element)

        footer_table = footer_type.add_table(rows=2, cols=1, width=Cm(footer_width_cm))
        # First row: thin light gray bar via cell shading
        # Ensure required imports are available in this scope
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        cell_bar = footer_table.cell(0, 0)
        shading_elm = parse_xml(r'<w:shd {} w:fill="CCCCCC"/>'.format(nsdecls('w')))  # Changed to light gray
        cell_bar._tc.get_or_add_tcPr().append(shading_elm)

        # Set the first row's height to 4px (approx 3pt)
        footer_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        footer_table.rows[0].height = Pt(3)  # ~4 pixels

        # Second row: text in Calibri, size 10, italics
        cell_text = footer_table.cell(1, 0)
        p = cell_text.paragraphs[0]
        run = p.add_run("Camtek Confidential and proprietary - distribution or any commercial or other external use of this document is forbidden.")
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Calculate available width for content
    available_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    text_width_cm = available_width_cm * 0.24  # Text column (now first)
    image_width_cm = available_width_cm * 0.76  # Image column (now second)
    
    
    # This helps mitigate the first page spacing issue by creating a fixed anchor point
    zero_table = document.add_table(rows=1, cols=1)
    # Use Table Grid style but we'll explicitly remove borders
    zero_table.style = 'Table Grid'
    zero_table.autofit = False
    zero_table.allow_autofit = False
    
    # Set to intended height of 1.70 cm - enforce exact height
    zero_table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    zero_table.rows[0].height = Cm(1.70)  # Changed from Pt(1.70) to Cm(1.70)
    
    # Make the cell completely transparent with no padding
    cell = zero_table.cell(0, 0)
    
    # Remove all borders at the cell level
    set_cell_border(
        cell,
        top={"sz": "0", "val": "none"},
        bottom={"sz": "0", "val": "none"},
        start={"sz": "0", "val": "none"},
        end={"sz": "0", "val": "none"}
    )
    
    # Remove cell margins/padding
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'+
                     f'<w:top w:w="0" w:type="dxa"/>'+
                     f'<w:bottom w:w="0" w:type="dxa"/>'+
                     f'<w:start w:w="0" w:type="dxa"/>'+
                     f'<w:end w:w="0" w:type="dxa"/>'+
                     f'</w:tcMar>')
    tcPr.append(tcMar)
    
    # Add explicit table border removal at the table level
    tbl_element = zero_table._element
    
    # Check if tblPr already exists
    tblPr = tbl_element.find('.//w:tblPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    
    # If not, create it
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl_element.insert(0, tblPr)
    
    # Create and append the borders element - explicitly set all borders to none
    tblBorders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'+
                        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'+
                        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'+
                        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'+
                        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'+
                        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'+
                        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'+
                        f'</w:tblBorders>')
    tblPr.append(tblBorders)
    
    # Force row height in XML to ensure it's applied - convert to twips (1/20 of a point)
    # For 1.70 cm, calculate equivalent in twips (1 cm ≈ 567 twips)
    twips_value = int(1.70 * 567)
    
    trPr = zero_table.rows[0]._tr.get_or_add_trPr()
    trHeight = parse_xml(f'<w:trHeight {nsdecls("w")} w:val="{twips_value}" w:hRule="exact"/>')
    
    # Remove any existing height settings
    for old_height in trPr.findall('.//w:trHeight', 
                               namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        trPr.remove(old_height)
        
    # Add new height setting
    trPr.append(trHeight)
    
    # Clear paragraph in cell and ensure it has no spacing
    p = cell.paragraphs[0]
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    
    # # Add side image to first page AFTER the zero-height table but BEFORE content
    # # --- Block commented out to remove side image ---
    # sideimage_path = os.path.join(assets_folder, "sideimage.png")
    # if os.path.exists(sideimage_path):
    #     # Create a special container for the side image that won't affect layout flow
    #     side_img_p = document.add_paragraph()
    #     side_img_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    #     side_img_p.paragraph_format.space_before = Pt(0)
    #     side_img_p.paragraph_format.space_after = Pt(0)
    #     run = side_img_p.add_run()
    #
    #     try:
    #         from PIL import Image
    #         with Image.open(sideimage_path) as img:
    #             aspect_ratio = img.width / img.height
    #
    #         height_cm = 10  # 10cm height
    #         width_cm = height_cm * aspect_ratio
    #
    #         run.add_picture(sideimage_path, width=Cm(width_cm), height=Cm(height_cm))
    #     except Exception as e:
    #         print(f"Warning: Could not add side image: {str(e)}")
    #         # Remove the paragraph if image addition failed
    #         if side_img_p:
    #             p_element = side_img_p._element
    #             p_element.getparent().remove(p_element)
    # # --- End of commented out block ---
    
    first_frame = True
    line_count = 0
    page_number = 1  # Track page number for different spacing needs
    
    for mapping in frame_step_mapping:
        if not os.path.exists(mapping['frame']):
            print(f"Warning: {mapping['frame']} does not exist.")
            continue
            
        # Skip frames without text
        if not mapping['text'].strip():
            continue
            
        if not first_frame:
            document.add_page_break()
            page_number += 1
            
        # For first page content, we completely eliminate ANY spacing elements
        if first_frame:
            # Directly create content table with no spacing
            adjusted_spacing = 0
        else:
            # Calculate spacing for non-first pages
            header_footer_buffer = 2.5
            available_height_cm = section.page_height.cm - section.top_margin.cm - section.bottom_margin.cm - header_footer_buffer
            content_height_cm = 10.0
            top_spacing_cm = max(1.0, (available_height_cm - content_height_cm) / 2.0)
            adjusted_spacing = top_spacing_cm
            # Add spacing using invisible table for non-first pages only
            add_spacing_table(document, adjusted_spacing)
            
        # Create main content table with absolutely minimal spacing
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Calculate widths in twips (1/20th of a point)
        # Convert the cm measurements to twips directly
        text_width_twips = int(text_width_cm * 567)  # 567 twips per cm
        image_width_twips = int(image_width_cm * 567)
        
        # Set table properties - enforce fixed layout
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}><w:tblLayout w:type="fixed"/><w:jc w:val="center"/></w:tblPr>')
        table._element.append(tblPr)
        
        # Set exact grid columns with precise widths (text first, then image)
        tblGrid = parse_xml(f'<w:tblGrid {nsdecls("w")}>' +
                          f'<w:gridCol w:w="{text_width_twips}"/>' +
                          f'<w:gridCol w:w="{image_width_twips}"/>' +
                          '</w:tblGrid>')
        table._element.append(tblGrid)
        
        # Apply strict fixed widths to each cell
        for i, width in enumerate([text_width_twips, image_width_twips]):
            cell = table.cell(0, i)._tc
            tcPr = cell.get_or_add_tcPr()
            
            # Remove any existing width setting
            for old_tcW in tcPr.findall('.//' + qn('w:tcW')):
                tcPr.remove(old_tcW)
                
            # Add new strict width setting
            tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width}" w:type="dxa"/>')
            tcPr.append(tcW)
            
            # Add noWrap property to prevent text from affecting table layout
            tcNoWrap = parse_xml(f'<w:noWrap {nsdecls("w")}/>')
            tcPr.append(tcNoWrap)
        
        # Convert twips back to cm for image sizing
        text_width_cm = text_width_twips / 567.0  # 567 twips per cm
        image_width_cm = image_width_twips / 567.0

        # Remove all possible spacing from table and cells
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO  # Allow row to shrink to content
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.space_before = Pt(0)
                    paragraph.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        # Remove all borders from the table
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"sz": "0", "val": "none"},
                    bottom={"sz": "0", "val": "none"},
                    start={"sz": "0", "val": "none"},
                    end={"sz": "0", "val": "none"}
                )
        
        # Set each cell's borders to white
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell,
                    start={"sz": "2", "val": "single", "color": "FFFFFF"},
                    end={"sz": "2", "val": "single", "color": "FFFFFF"},
                    top={"sz": "2", "val": "single", "color": "FFFFFF"},
                    bottom={"sz": "2", "val": "single", "color": "FFFFFF"},
                )
        
        # Text cell - configure first (now on the left)
        cell_text = table.cell(0, 0)
        cell_text.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        
        # Apply a right margin to separate text from image
        from docx.oxml.ns import nsdecls  # Ensure nsdecls is imported in this scope
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:right w:w="200" w:type="dxa"/></w:tcMar>')
        cell_text._tc.get_or_add_tcPr().append(tcMar)
        
        # Clear any existing paragraphs
        for p in cell_text.paragraphs[1:]:
            p._element.getparent().remove(p._element)
        p_text = cell_text.paragraphs[0]
        
        # Add vertical border to text cell with the requested dark teal color (now on right side)
        set_cell_border(cell_text,
            end={"sz": "20", "val": "single", "color": "008760"})

        # Image cell - configure second (now on the right)
        cell_image = table.cell(0, 1)
        cell_image.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Clear any existing content in the image cell
        for p in cell_image.paragraphs[1:]:
            p._element.getparent().remove(p._element)
        p_img = cell_image.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ensure proper vertical spacing within the cell (balanced top/bottom)
        p_img.space_before = Pt(6)
        p_img.space_after = Pt(6)
        r_img = p_img.add_run()

        # Add image with precise sizing
        try:
            # Get image dimensions to maintain aspect ratio
            from PIL import Image
            with Image.open(mapping['frame']) as img:
                img_width, img_height = img.size
                aspect_ratio = img_width / img_height
                
            # Calculate optimal image size to fit while maintaining aspect ratio
            max_width_cm = image_width_cm * 0.98  # Use 98% of cell width (slight margin)
            max_height_cm = 12.5  # Increased from 10.0 to allow taller images
            
            # Determine if width or height is the limiting factor
            proposed_height = max_width_cm / aspect_ratio
            if proposed_height > max_height_cm:
                # Height limited, calculate width based on max height
                final_width = max_height_cm * aspect_ratio
            else:
                # Width limited
                final_width = max_width_cm
                
            r_img.add_picture(mapping['frame'], width=Cm(final_width))
        except Exception as e:
            print(f"Warning: Could not add image {mapping['frame']}: {str(e)}")
            continue
            
        # Add numbered steps with proper formatting
        lines = mapping['text'].split('\n')
        for i, line_text in enumerate(lines):
            if not line_text.strip():
                continue
            line_count += 1
            
            if i > 0:
                p_text = cell_text.add_paragraph()
            
            # Add step number
            step_run = p_text.add_run(f"{line_count}. ")
            step_run.font.name = 'Arial'
            step_run.font.size = Pt(14)
            step_run.font.bold = True
            
            # Add step text
            text_run = p_text.add_run(line_text)
            text_run.font.name = 'Arial'
            text_run.font.size = Pt(12)
            text_run.font.bold = False
        
        first_frame = False
    
    # Add copyright page
    document.add_page_break()
    
    # Read and copy content from Copyright.docx
    copyright_path = os.path.join(assets_folder, "Copyright.docx")
    copyright_doc = Document(copyright_path)
    
    # Copy all content including paragraphs AND tables from Copyright.docx
    for element in copyright_doc.element.body:
        # Check if this is a paragraph
        if element.tag.endswith('p'):
            paragraph = [p for p in copyright_doc.paragraphs if p._element is element][0]
            
            # Create new paragraph in target document
            new_para = document.add_paragraph()
            
            # Copy paragraph formatting
            new_para.paragraph_format._element = paragraph.paragraph_format._element
            
            # Copy runs with their formatting
            for run in paragraph.runs:
                new_run = new_para.add_run(run.text)
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic
                new_run.font.underline = run.font.underline
                
            # Copy paragraph alignment
            new_para.alignment = paragraph.alignment
            
        # Check if this is a table
        elif element.tag.endswith('tbl'):
            # Find the corresponding table object
            table_index = copyright_doc.element.body.index(element)
            table_count = 0
            original_table = None
            
            # Find which table this is in the document
            for t in copyright_doc.tables:
                if t._element is element:
                    original_table = t
                    break
                table_count += 1
            
            if original_table:
                # Create a new table with same dimensions
                rows = len(original_table.rows)
                cols = len(original_table.rows[0].cells) if rows > 0 else 1
                new_table = document.add_table(rows=rows, cols=cols)
                
                # Copy table style
                new_table.style = original_table.style
                
                # Copy table properties from XML
                # Get table properties from original table
                if hasattr(original_table._element, "tblPr") and original_table._element.tblPr is not None:
                    tblPr = parse_xml(original_table._element.tblPr.xml)
                    new_table._element.append(tblPr)
                
                # Copy content and formatting of each cell
                for i, row in enumerate(original_table.rows):
                    # Set row height
                    if row.height_rule:
                        new_table.rows[i].height_rule = row.height_rule
                    if row.height:
                        new_table.rows[i].height = row.height
                    
                    for j, cell in enumerate(row.cells):
                        # Copy cell content and formatting
                        new_cell = new_table.cell(i, j)
                        
                        # Remove default paragraph in new cell
                        for p in new_cell.paragraphs:
                            p._element.getparent().remove(p._element)
                        
                        # Copy each paragraph from original cell
                        for paragraph in cell.paragraphs:
                            new_para = new_cell.add_paragraph()
                            
                            # Copy paragraph text and formatting
                            for run in paragraph.runs:
                                new_run = new_para.add_run(run.text)
                                new_run.font.name = run.font.name
                                new_run.font.size = run.font.size
                                new_run.font.bold = run.font.bold
                                new_run.font.italic = run.font.italic
                                new_run.font.underline = run.font.underline
                            
                            # Copy paragraph alignment
    # Save the document
    document.save(output_path)
    
    # Return only the Word document path (no PDF conversion)
    return output_path

def convert_docx_to_pdf(docx_file, pdf_file):
    """
    Convert Word document to PDF
    
    For PythonAnywhere, we'll use a different approach since win32com is not available
    This function will be modified to use a different conversion method
    """
    try:
        # Try to use win32com if available (Windows only)
        import win32com.client
        import time
        
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        
        # Get absolute paths
        docx_path = os.path.abspath(docx_file)
        
        # Ensure the source file exists
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Word document not found: {docx_path}")
            
        # Open document with read-only and no alert flags
        doc = word.Documents.Open(
            docx_path,
            ReadOnly=True,
            Visible=False,
            ConfirmConversions=False
        )
        
        # Wait a moment for document to fully load
        time.sleep(2)
        
        # Save as PDF
        try:
            doc.ExportAsFixedFormat(
                OutputFileName=pdf_file,
                ExportFormat=17,  # wdExportFormatPDF
                OpenAfterExport=False,
                OptimizeFor=0,    # wdExportOptimizeForPrint
                Range=0,          # wdExportAllDocument
                IncludeDocProps=True,
                KeepIRM=True, 
                DocStructureTags=True,
                BitmapMissingFonts=True,
                UseISO19005_1=False
            )
        except Exception as save_error:
            # Alternative method if ExportAsFixedFormat fails
            doc.SaveAs(pdf_file, FileFormat=17)
        
        # Properly close the document
        doc.Close(SaveChanges=False)
        
        # Clean up
        word.Quit()
        
    except ImportError:
        # For PythonAnywhere (Linux), use alternative conversion
        try:
            # Try using LibreOffice if available
            subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', os.path.dirname(pdf_file), docx_file
            ], check=True)
        except (subprocess.SubprocessError, FileNotFoundError):
            # If LibreOffice fails, try using unoconv
            try:
                subprocess.run(['unoconv', '-f', 'pdf', '-o', pdf_file, docx_file], check=True)
            except (subprocess.SubprocessError, FileNotFoundError):
                # If all conversion methods fail, raise an error
                raise Exception("PDF conversion failed. Neither win32com, LibreOffice, nor unoconv are available.")