import http.server
import socketserver
import os
import json
import sys
import tempfile
from docx import Document
from docx.shared import RGBColor
from openai import OpenAI
from config import OPENAI_API_KEY

# Initialize OpenAI client (v1+)
openai_client = OpenAI(api_key=OPENAI_API_KEY)
import shutil
from urllib.parse import parse_qs, urlparse
import mimetypes
import threading
from pathlib import Path
import io
import email
from werkzeug.utils import secure_filename # Added import
import websockets
import asyncio
import keyboard
import pyperclip
import openpyxl
from openpyxl import Workbook
import time
from datetime import datetime
import win32com.client
import pythoncom
import tkinter as tk
from tkinter import filedialog

# Import functions from the formatting script
from formatting import clean_text, remove_blank_rows, process_excel

# Import functions from Video to PDF project
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'Video to PDF'))
from final_app import process_files

# Import functions from TTS project
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'TTS'))
# Assuming tts4.py has functions get_available_voices() and generate_tts()
try:
    from tts4 import get_available_voices, generate_tts
    TTS_ENABLED = True
except ImportError:
    print("WARNING: TTS module (tts4.py) not found or failed to import. TTS functionality will be disabled.")
    TTS_ENABLED = False
    # Define dummy functions if import fails
    def get_available_voices(): return []
    def generate_tts(text, voice_id): raise NotImplementedError("TTS module not loaded")

def handle_error(e):
    """Simple error handler that returns a user-friendly error message"""
    error_msg = str(e)
    if "google.cloud" in error_msg.lower():
        return "Error with Google Cloud services. Please check your credentials."
    elif "ffmpeg" in error_msg.lower():
        return "Error processing video. Please check if the video file is valid."
    elif "excel" in error_msg.lower():
        return "Error processing Excel file. Please check if the file is valid."
    else:
        return f"An error occurred: {error_msg}"

# Define the ports
HTTP_PORT = 8001
WS_PORT = 8002

# Get the directory of the current script
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(SCRIPT_DIR, 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Store active Excel sessions
excel_sessions = {}

class ExcelSession:
    def __init__(self, filename, save_location):
        self.filename = filename
        self.save_location = save_location
        self.workbook = Workbook()
        self.sheet = self.workbook.active
        self.row = 1
        self.websocket = None
        self.monitoring = False

    async def add_text(self, text):
        # Add the text to Excel
        self.sheet.cell(row=self.row, column=1).value = text
        
        # Add a blank row
        self.row += 1
        self.sheet.cell(row=self.row, column=1).value = ''
        
        # Move to next row
        self.row += 1

        # Send update to client
        if self.websocket:
            await self.websocket.send(json.dumps({
                'type': 'update',
                'text': text,
                'row': self.row - 2
            }))

    def save(self):
        save_path = os.path.join(self.save_location, f"{self.filename}.xlsx")
        self.workbook.save(save_path)
        return save_path

class MultipartFormParser:
    def __init__(self, content_type, content_length, rfile):
        self.content_type = content_type
        self.content_length = content_length
        self.rfile = rfile
        print(f"MultipartFormParser initialized with content length: {content_length}")  # Debug log
        
    def parse(self):
        print("Starting form parse")  # Debug log
        fields = {}

        try:
            # Get boundary from content type
            boundary = self.content_type.split('boundary=')[1].encode()
            print(f"Found boundary: {boundary}")  # Debug log
            remainbytes = self.content_length
            
            # Read until first boundary
            while True:
                line = self.rfile.readline()
                remainbytes -= len(line)
                if boundary in line:
                    print("Found first boundary")  # Debug log
                    break

            # Process form fields
            while remainbytes > 0:
                try:
            # Parse headers
                    # Parse headers
                    headers = {}
                    while True:
                        line = self.rfile.readline()
                        remainbytes -= len(line)
                        if line == b'\r\n':
                            break
                        
                        # Parse header line
                        line = line.decode('utf-8').strip()
                        if ':' in line:
                            key, value = line.split(':', 1)
                            headers[key.strip().lower()] = value.strip()
                    
                    # Get field name and filename from Content-Disposition
                    filename = None
                    name = None
                    if 'content-disposition' in headers:
                        disposition = headers['content-disposition']
                        items = disposition.split(';')
                        for item in items:
                            item = item.strip()
                            if item.startswith('name='):
                                name = item.split('=')[1].strip('"')
                            elif item.startswith('filename='):
                                filename = item.split('=')[1].strip('"')
                        if not name:  # If name wasn't found, skip this part
                            continue
                    else:
                        continue
                    
                    # Read content until boundary
                    content = io.BytesIO()
                    while remainbytes > 0:
                        line = self.rfile.readline()
                        remainbytes -= len(line)
                        if boundary in line:
                            break
                        content.write(line)
                    
                    # Store field value and filename
                    value = content.getvalue().strip(b'\r\n')
                    if filename:
                        print(f"Found file: {filename}, content length: {len(value)}")  # Debug log
                        # Store as a tuple: (filename, content)
                        fields[name.encode()] = (filename.encode(), value)
                    else:
                        print(f"Found field: {name}")  # Debug log
                        # Store just the value
                        fields[name.encode()] = value

                    if remainbytes <= 0:
                        break

                except Exception as e:
                    print(f"Error processing field: {str(e)}")  # Debug log
                    continue
        except Exception as e:
            print(f"Error in main form processing loop: {str(e)}")  # Debug log
            return fields
        
        print(f"Form parsing complete. Found fields: {list(fields.keys())}")  # Debug log
        return fields

class MultiToolHandler(http.server.SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=SCRIPT_DIR, **kwargs)
    
    def do_GET(self):
        parsed_path = urlparse(self.path)
        path = parsed_path.path
        
        if path == '/':
            self.path = '/index.html'
        elif path == '/shutdown':
            self.send_response(200)
            self.send_header('Content-type', 'text/html')
            self.end_headers()
            self.wfile.write(b'Server shutting down...')
            threading.Thread(target=self.server.shutdown).start()
            return
        elif path == '/get_frames':
            self.redirect_to('/templates/get_frames.html')
            return
        elif path == '/convert_to_pdf':
            self.redirect_to('/templates/convert_pdf.html')
            return
        elif path == '/settings':
            self.redirect_to('/templates/settings.html')
            return
        elif path.startswith('/ws-port'):
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'port': WS_PORT}).encode())
            return
        elif path == '/api/browse-directory':
            self.handle_browse_directory()
            return
        elif path == '/api/history':
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            history = self.get_upload_history()
            self.wfile.write(json.dumps(history).encode())
            return
        elif path == '/api/tts-voices':
            self.handle_get_tts_voices()
            return
        elif path.startswith('/download/'):
            # Extract job_id and filename from path
            parts = path.split('/')
            if len(parts) >= 4:
                job_id = parts[2]
                filename = parts[3]
                job_dir = os.path.join(UPLOAD_DIR, job_id)
                
                # Try finding the file directly in job_dir (likely Word-to-PDF)
                file_path_direct = os.path.join(job_dir, filename)
                file_path = None

                if os.path.exists(file_path_direct):
                    file_path = file_path_direct
                else:
                    # If not found directly, check in 'results' subdir (likely Video-to-PDF)
                    results_dir = os.path.join(job_dir, 'results')
                    file_path_results = os.path.join(results_dir, filename)
                    if os.path.exists(file_path_results):
                        file_path = file_path_results

                if file_path and os.path.exists(file_path):
                    with open(file_path, 'rb') as f:
                        self.send_response(200)
                        if filename.endswith('.pdf'):
                            self.send_header('Content-type', 'application/pdf')
                        elif filename.endswith('.docx'):
                            self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
                        self.end_headers()
                        shutil.copyfileobj(f, self.wfile)
                    return # Correct indentation for the return statement
        elif os.path.exists(path[1:]):  # Remove leading slash
            with open(path[1:], 'rb') as f:
                self.send_response(200)
                if path.endswith('.xlsx'):
                    self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                self.end_headers()
                shutil.copyfileobj(f, self.wfile)
            return
        
        return super().do_GET()
    
    def do_POST(self):
        if self.path == '/api/format-excel':
            self.handle_format_excel()
        elif self.path == '/upload':
            self.handle_video_upload()
        elif self.path.startswith('/process/'):
            self.handle_video_processing()
        elif self.path.startswith('/cleanup/'):
            self.handle_cleanup()
        elif self.path == '/api/start-excel-session':
            self.handle_start_excel_session()
        elif self.path == '/api/convert-word':
            self.handle_word_conversion()
        elif self.path == '/api/process-checklist':
            if not OPENAI_API_KEY or OPENAI_API_KEY == 'your-api-key-here':
                self.send_error(500, "OpenAI API key not configured. Please set OPENAI_API_KEY in config.py")
                return
            try:
                self.handle_checklist_processing()
            except Exception as e:
                print(f"Error in checklist processing: {str(e)}")  # Debug log
                self.send_error(500, str(e))
        elif self.path == '/api/tts':
            self.handle_tts_conversion()

    def handle_checklist_processing(self):
        """Handles POST request for processing Word documents into checklists."""
        print("Starting checklist processing...")  # Debug log
        try:
            print(f"Headers: {self.headers}")  # Debug log
            content_length = int(self.headers['Content-Length'])
            content_type = self.headers['Content-Type']
            
            if not content_type.startswith('multipart/form-data'):
                raise ValueError("Expected multipart/form-data")
            
            parser = MultipartFormParser(content_type, content_length, self.rfile)
            form = parser.parse()
            print(f"Form fields: {list(form.keys())}")  # Debug log
            
            if not form or b'file' not in form:
                print("No file found in form data")  # Debug log
                raise ValueError("No file uploaded")

            # Extract filename and content
            filename_bytes, file_content = form[b'file']
            original_filename = filename_bytes.decode('utf-8', errors='ignore')
            print(f"Original filename: {original_filename}")  # Debug log
            original_filename_safe = secure_filename(original_filename)
            base_name, _ = os.path.splitext(original_filename_safe)
            
            # Create a unique ID for this conversion
            conversion_id = os.urandom(16).hex()
            conversion_dir = os.path.join(UPLOAD_DIR, conversion_id)
            os.makedirs(conversion_dir, exist_ok=True)

            # Save the uploaded file
            input_path = os.path.join(conversion_dir, original_filename_safe)
            with open(input_path, 'wb') as f:
                f.write(file_content)

            # Read the Word document
            doc = Document(input_path)
            text_content = []
            found_first_step = False
            
            # Extract text content after finding the first step
            # First, collect all paragraphs
            all_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # Try to find where the steps begin
            for i, text in enumerate(all_paragraphs):
                # Look for common step indicators
                if (text.lower().startswith(('step', 'procedure', 'assembly')) or
                    (any(char.isdigit() for char in text[:5]) and len(text) > 10) or  # Number in first 5 chars
                    text.lower().startswith(('1.', '1)', '1 -', '[1]', '(1)'))):  # Common numbering formats
                    found_first_step = True
                    text_content = all_paragraphs[i:]  # Include this and all following paragraphs
                    break
            
            # If no clear step indicator found, include all non-empty paragraphs
            if not found_first_step and all_paragraphs:
                text_content = all_paragraphs

            # Prepare prompt for OpenAI
            prompt = """Analyze the following assembly or procedure text and create a structured checklist starting from step 6 onwards.
            Format the output as a table with one column:
            Column 1: Step/Sub-step number followed by a short summary of the step (e.g., "6. Install mounting brackets")
            
            Guidelines:
            - ONLY include steps and substeps from step 6 onwards
            - IMPORTANT: Break down ALL steps (not just 6.1) into detailed sub-steps
            - For EVERY main step (6, 7, 8, etc.), provide detailed substeps (6.1, 6.2, 7.1, 7.2, etc.)
            - For complex substeps, provide further breakdown (6.1.1, 6.1.2, 6.2.1, etc.)
            - Create concise summaries that capture the essence of each step
            - IMPORTANT: Include part numbers in the summaries wherever they are mentioned
            - Preserve all technical details and specifications in the summary
            - Use clear, concise language
            - Maintain the original sequence of operations
            - Include any critical warnings or notes as sub-steps
            - Ensure each step and substep is on its own separate line
            - Use clear numbering (6., 6.1., 6.1.1, 6.2., 7., 7.1., etc.) for all steps and substeps
            - Do not include any table formatting characters like |, just the step number and summary
            - Make sure to break down EVERY step, not just step 6.1

            Text to process:
            """ + "\n".join(text_content)

            # Call OpenAI API
            response = openai_client.chat.completions.create(
                model="gpt-4.1",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that extracts and organizes steps from text into a structured table format."},
                    {"role": "user", "content": prompt}
                ]
            )

            # Create new Word document
            output_doc = Document()
            
            # Set document properties for standard page formatting with integer values
            section = output_doc.sections[0]
            # Standard letter size (using integer values)
            section.page_width = int(8.5 * 1440)  # 8.5 inches in twips (1440 twips per inch)
            section.page_height = int(11 * 1440)  # 11 inches in twips
            # Normal margins (using integer values)
            section.left_margin = section.right_margin = int(1 * 1440)  # 1 inch margins
            
            # Add title with proper formatting
            title = output_doc.add_heading('Assembly Checklist - Steps 6 and Beyond', 0)
            title.alignment = 1  # Center alignment
            
            # Add introduction with proper formatting
            intro1 = output_doc.add_paragraph()
            intro1.add_run(f"Checklist generated from: {original_filename}").bold = True
            intro1.alignment = 0  # Left alignment
            
            intro2 = output_doc.add_paragraph("This checklist contains steps 6 and beyond from the assembly procedure. Follow each step carefully and complete the values and pass/fail columns as appropriate.")
            intro2.alignment = 0  # Left alignment
            
            intro3 = output_doc.add_paragraph()
            intro3.add_run("Note: Part numbers are included in step descriptions where applicable.").italic = True
            intro3.alignment = 0  # Left alignment
            
            # Add spacing before table
            output_doc.add_paragraph()
            
            # Parse the OpenAI response and create a table
            table_content = response.choices[0].message.content
            
            # Split the content into lines and filter out empty lines
            lines = [line.strip() for line in table_content.split('\n') if line.strip()]
            
            # Create table with appropriate number of rows
            table = output_doc.add_table(rows=len(lines), cols=3)
            table.style = 'Table Grid'
            
            # Set table width to 100% of available width between margins
            table.autofit = False
            table.allow_autofit = False
            
            # Set preferred widths for columns using integer values
            table.columns[0].width = int(4.9 * 1440)  # Step number and summary column (75% of available width)
            table.columns[1].width = int(1.0 * 1440)  # Values column (15% of available width)
            table.columns[2].width = int(0.6 * 1440)  # Pass/Fail column (10% of available width)
            
            # Add headers with proper formatting
            header_cells = table.rows[0].cells
            header_cells[0].text = "Step and Description"
            header_cells[1].text = "Values"
            header_cells[2].text = "Pass/Fail"
            
            # Format header row
            for cell in header_cells:
                # Bold headers
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                # Center align headers
                cell.paragraphs[0].alignment = 1  # Center alignment
                # Note: Cell shading removed as python-docx doesn't support cell.fill directly
            
            # Add content
            for i, line in enumerate(lines[1:], start=1):  # Skip header row
                # The line should already contain step number and summary
                step_summary = line.strip()
                
                # Add to table
                row_cells = table.rows[i].cells
                row_cells[0].text = step_summary  # Step number and summary
                row_cells[1].text = ""  # Values column (blank)
                row_cells[2].text = ""  # Pass/Fail column
                
                # Format cells properly
                for cell in row_cells:
                    # Set vertical alignment to center for all cells
                    cell.vertical_alignment = 1  # 1 = CENTER
                    
                    # Ensure proper paragraph formatting
                    for paragraph in cell.paragraphs:
                        # Add minimal spacing for better readability (using integer values)
                        paragraph.paragraph_format.space_before = 24  # 1/60 of an inch (integer)
                        paragraph.paragraph_format.space_after = 24   # 1/60 of an inch (integer)
                        # Line spacing - use line_spacing_rule instead of line_spacing to avoid float issues
                        paragraph.paragraph_format.line_spacing_rule = 1  # Single spacing (1=single, 2=double)
                
                # Process step numbers for better alignment
                if '.' in step_summary[:10]:  # Extended search range to catch multi-level substeps
                    # Find the position of the first period
                    dot_pos = step_summary.find('.')
                    # Check if there's content after the period
                    if dot_pos > 0 and dot_pos < len(step_summary) - 1:
                        # Extract everything before the first period as the main step number
                        main_step = step_summary[:dot_pos].strip()
                        
                        # Find the position of the first space after the step number
                        # This helps separate the step number from the summary text
                        space_pos = step_summary.find(' ', dot_pos)
                        if space_pos > 0:
                            # Extract the full step number (including all substep levels)
                            full_step_num = step_summary[:space_pos].strip()
                            # Extract the summary text
                            summary = step_summary[space_pos:].strip()
                            
                            # Check if this is a main step or sub-step
                            try:
                                # Try to convert main_step to int to check if it's a main step
                                int(main_step)
                                is_main_step = len(main_step) == 1  # Main steps are single digits (1, 2, etc.)
                            except ValueError:
                                # If conversion fails, it's not a standard step number
                                is_main_step = False
                            
                            # Format based on step level
                            if is_main_step:
                                # Format main steps
                                row_cells[0].paragraphs[0].alignment = 0  # Left align
                                p = row_cells[0].paragraphs[0].clear()
                                p.add_run(f"{full_step_num} ").bold = True
                                p.add_run(summary)
                            else:
                                # Format sub-steps with indentation
                                row_cells[0].paragraphs[0].alignment = 0  # Left align
                                row_cells[0].paragraphs[0].paragraph_format.left_indent = int(0.25 * 1440)  # 0.25 inches (as integer)
                                p = row_cells[0].paragraphs[0].clear()
                                p.add_run(f"{full_step_num} ").bold = True
                                p.add_run(summary)
            
            # Add footer with proper formatting
            output_doc.add_paragraph()  # Add spacing
            notes_heading = output_doc.add_paragraph()
            notes_heading.add_run("Notes:").bold = True
            
            # Add lines for notes
            for _ in range(3):
                line = output_doc.add_paragraph()
                line.add_run("_" * 80)
                line.paragraph_format.space_after = int(1/6 * 1440)  # Add space between lines (1/6 inch as integer)

            # Save the output document
            output_filename = f"{base_name}_checklist.docx"
            output_path = os.path.join(conversion_dir, output_filename)
            output_doc.save(output_path)

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({
                'id': conversion_id,
                'filename': output_filename,
                'success': True
            }).encode())

        except Exception as e:
            print(f"Error in checklist processing: {str(e)}")  # Debug log
            self.send_error(500, str(e))

    def handle_word_conversion(self):
        try:
            content_length = int(self.headers['Content-Length'])
            content_type = self.headers['Content-Type']
            
            if not content_type.startswith('multipart/form-data'):
                raise ValueError("Expected multipart/form-data")
            
            parser = MultipartFormParser(content_type, content_length, self.rfile)
            form = parser.parse()

            if not form or b'file' not in form:
                raise ValueError("No file uploaded")

            # Extract original filename and content
            word_filename_bytes, word_content = form[b'file']
            original_filename = word_filename_bytes.decode('utf-8', errors='ignore')
            original_filename_safe = secure_filename(original_filename)
            base_name, _ = os.path.splitext(original_filename_safe)
            pdf_filename_safe = f"{base_name}.pdf"

            # Create a unique ID for this conversion
            conversion_id = os.urandom(16).hex()
            conversion_dir = os.path.join(UPLOAD_DIR, conversion_id)
            os.makedirs(conversion_dir, exist_ok=True)

            # Save the Word file using original (sanitized) name
            word_path = os.path.join(conversion_dir, original_filename_safe)
            with open(word_path, 'wb') as f:
                f.write(word_content)

            # Convert to PDF using original (sanitized) base name
            pdf_path = os.path.join(conversion_dir, pdf_filename_safe)
            self.convert_word_to_pdf(word_path, pdf_path)

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({
                'id': conversion_id,
                'pdf_filename': pdf_filename_safe, # Return the actual PDF filename
                'success': True
            }).encode())

        except Exception as e:
            self.send_error(500, str(e))

    def convert_word_to_pdf(self, word_path, pdf_path):
        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch('Word.Application')
            doc = word.Documents.Open(word_path)
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 = PDF format
            doc.Close()
            word.Quit()
        finally:
            pythoncom.CoUninitialize()

    def redirect_to(self, path):
        self.send_response(302)
        self.send_header('Location', path)
        self.end_headers()
    
    def handle_format_excel(self):
        content_length = int(self.headers['Content-Length'])
        content_type = self.headers['Content-Type']
        
        if not content_type.startswith('multipart/form-data'):
            self.send_error(400, "Expected multipart/form-data")
            return
        
        parser = MultipartFormParser(content_type, content_length, self.rfile)
        form = parser.parse()
        
        if not form or b'file' not in form:
            self.send_error(400, "No file uploaded")
            return
        
        file_tuple = form[b'file']
        
        if not file_tuple or not isinstance(file_tuple, tuple):
            self.send_error(400, "Not a file")
            return
            
        filename, file_content = file_tuple
        
        if not file_content:
            self.send_error(400, "Not a file")
            return
        
        options = {
            'remove_blank_lines': form.get(b'removeBlankLines', b'false').decode() == 'true',
            'capitalize_sentences': form.get(b'capitalizeSentences', b'false').decode() == 'true',
            'add_periods': form.get(b'addPeriods', b'false').decode() == 'true',
            'remove_spaces_quotes': form.get(b'removeSpacesQuotes', b'false').decode() == 'true',
            'remove_spaces_unquoted': form.get(b'removeSpacesUnquoted', b'false').decode() == 'true',
            'remove_lone_quotes': form.get(b'removeLoneQuotes', b'false').decode() == 'true',
            'remove_ellipsis': form.get(b'removeEllipsis', b'false').decode() == 'true'
        }
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            tmp_file.write(file_content)
            tmp_filename = tmp_file.name
        
        try:
            print(f"Starting formatting for {tmp_filename} with options: {options}") # DEBUG
            # First process the text if any text processing options are enabled
            text_options_enabled = any([options['capitalize_sentences'], options['add_periods'],
                                        options['remove_spaces_quotes'], options['remove_spaces_unquoted'],
                                        options['remove_lone_quotes'], options['remove_ellipsis']])
            if text_options_enabled:
                print(f"Calling process_excel for {tmp_filename}...") # DEBUG
                process_excel(tmp_filename, options) # Pass the options dict
                print(f"Finished process_excel for {tmp_filename}.") # DEBUG
            else:
                 print(f"Skipping process_excel for {tmp_filename} as no text options enabled.") # DEBUG

            # Then remove blank rows if that option is enabled
            if options['remove_blank_lines']:
                print(f"Calling remove_blank_rows for {tmp_filename}...") # DEBUG
                remove_blank_rows(tmp_filename)
                print(f"Finished remove_blank_rows for {tmp_filename}.") # DEBUG
            else:
                 print(f"Skipping remove_blank_rows for {tmp_filename}.") # DEBUG

            print(f"Reading processed data from {tmp_filename}...") # DEBUG
            with open(tmp_filename, 'rb') as f:
                processed_data = f.read()
            
            self.send_response(200)
            self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            self.send_header('Content-Disposition', 'attachment; filename="modified_file.xlsx"')
            self.send_header('Content-Length', str(len(processed_data)))
            self.end_headers()
            self.wfile.write(processed_data)
            
        except Exception as e:
            print(f"ERROR during Excel processing: {e}") # DEBUG - Print error to console
            self.send_error(500, f"Error processing Excel: {str(e)}")

        finally:
            if os.path.exists(tmp_filename):
                os.unlink(tmp_filename)

    def handle_video_upload(self):
        try:
            content_length = int(self.headers['Content-Length'])
            content_type = self.headers['Content-Type']
            
            if not content_type.startswith('multipart/form-data'):
                raise ValueError("Expected multipart/form-data")
            
            parser = MultipartFormParser(content_type, content_length, self.rfile)
            form = parser.parse()
            
            if not form or b'video' not in form or b'excel' not in form:
                raise ValueError("Missing video or excel file")

            # Extract filenames and content from the form data
            video_filename_bytes, video_content = form[b'video']
            excel_filename_bytes, excel_content = form[b'excel']

            # Decode filenames
            video_filename = video_filename_bytes.decode('utf-8', errors='ignore')
            excel_filename = excel_filename_bytes.decode('utf-8', errors='ignore')

            # Sanitize filenames (optional but recommended)
            video_filename_safe = secure_filename(video_filename)
            excel_filename_safe = secure_filename(excel_filename)

            job_id = os.urandom(16).hex()
            job_dir = os.path.join(UPLOAD_DIR, job_id)
            os.makedirs(job_dir)

            # Use original (sanitized) filenames
            video_path = os.path.join(job_dir, video_filename_safe)
            excel_path = os.path.join(job_dir, excel_filename_safe)

            with open(video_path, 'wb') as f:
                f.write(video_content)
            with open(excel_path, 'wb') as f:
                f.write(excel_content)

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'job_id': job_id}).encode())

        except Exception as e:
            self.send_error(500, str(e))

    def handle_video_processing(self):
        try:
            job_id = self.path.split('/')[-1]
            job_dir = os.path.join(UPLOAD_DIR, job_id)
            
            if not os.path.exists(job_dir):
                raise ValueError("Invalid job ID")

            # Find the video and excel files in the job directory
            video_path = None
            excel_path = None
            for filename in os.listdir(job_dir):
                if filename.lower().endswith(('.mp4', '.avi', '.mov', '.wmv', '.flv')): # Add other video extensions if needed
                    video_path = os.path.join(job_dir, filename)
                elif filename.lower().endswith(('.xlsx', '.xls')):
                    excel_path = os.path.join(job_dir, filename)

            if not video_path or not excel_path:
                 raise ValueError(f"Missing required video or excel file in job directory: {job_dir}")
            if not os.path.exists(video_path) or not os.path.exists(excel_path):
                 raise ValueError(f"Required files not found: Video={video_path}, Excel={excel_path}")

            # Create frames folder
            frames_folder = os.path.join(job_dir, 'frames')
            os.makedirs(frames_folder, exist_ok=True)

            # Create required folders
            temp_folder = os.path.join(job_dir, 'temp')
            # frames_folder is already defined and created
            results_folder = os.path.join(job_dir, 'results')
            os.makedirs(temp_folder, exist_ok=True)
            # os.makedirs(frames_folder, exist_ok=True) # Redundant
            os.makedirs(results_folder, exist_ok=True)

            # Copy required assets to results directory
            video_to_pdf_dir = os.path.join(SCRIPT_DIR, 'Video to PDF')
            assets = {
                'logo.jpg': 'logo.jpg',
                'sideimage.png': 'sideimage.png.png',  # Fix double extension
                'Copyright.docx': 'Copyright.docx',
                'Template.docx': 'Template.docx'
            }
            for dest_name, src_name in assets.items():
                src = os.path.join(video_to_pdf_dir, src_name)
                dst = os.path.join(results_folder, dest_name)
                if os.path.exists(src):
                    shutil.copy(src, dst)
                else:
                    raise ValueError(f"Required asset not found: {src}")

            # Process files using the Video to PDF project's function
            result = process_files(
                video_path=video_path,
                excel_path=excel_path,
                credential_path=os.path.join(UPLOAD_DIR, 'sapheb-b87c6918d4ef.json'),
                bucket_name='sap_pdf',
                frames_folder=frames_folder,
                temp_folder=temp_folder,
                job_id=job_id
            )

            # Ensure the docx file exists and get the base name for PDF
            results_dir = os.path.join(job_dir, 'results')
            docx_filename = result['docx_filename']
            docx_path = os.path.join(results_dir, docx_filename)
            if not os.path.exists(docx_path):
                raise ValueError(f"Generated DOCX file not found: {docx_path}")

            # Construct the original PDF path (as generated by processing.py)
            original_pdf_filename = docx_filename.replace('.docx', '.pdf')
            original_pdf_path = os.path.join(results_dir, original_pdf_filename)
            if not os.path.exists(original_pdf_path):
                 raise ValueError(f"Generated PDF file not found: {original_pdf_path}")

            # --- Apply new filename formatting ---
            # Get base name without '-PDF.pdf'
            base_name_with_underscores = original_pdf_filename.replace('-PDF.pdf', '')
            # Replace underscores with spaces
            base_name_with_spaces = base_name_with_underscores.replace('_', ' ')
            
            # Find the first space
            first_space_index = base_name_with_spaces.find(' ')
            
            if first_space_index != -1:
                # Construct new name: [FirstPart] Rest Of Name.pdf
                first_part = base_name_with_spaces[:first_space_index]
                rest_of_name = base_name_with_spaces[first_space_index:] # Includes the leading space
                new_pdf_filename = f"[{first_part}]{rest_of_name}.pdf"
            else:
                # No spaces, just add brackets: [WholeName].pdf
                new_pdf_filename = f"[{base_name_with_spaces}].pdf"

            # Sanitize the new filename
            new_pdf_filename_safe = secure_filename(new_pdf_filename)
            new_pdf_path = os.path.join(results_dir, new_pdf_filename_safe)

            # Rename the generated PDF file
            try:
                os.rename(original_pdf_path, new_pdf_path)
            except OSError as rename_error:
                 raise OSError(f"Failed to rename PDF: {rename_error}")
            # --- End of filename formatting ---

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({
                 # Send the *new* filename for download
                'download_url': f'/download/{job_id}/{new_pdf_filename_safe}'
            }).encode())

        except Exception as e:
            error_message = handle_error(e)
            self.send_error(500, error_message)

    def handle_cleanup(self):
        try:
            job_id = self.path.split('/')[-1]
            job_dir = os.path.join(UPLOAD_DIR, job_id)
            
            if os.path.exists(job_dir):
                shutil.rmtree(job_dir)
            
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'success': True}).encode())

        except Exception as e:
            self.send_error(500, str(e))

    def handle_start_excel_session(self):
        try:
            content_length = int(self.headers['Content-Length'])
            data = self.rfile.read(content_length)
            params = json.loads(data)
            
            session_id = os.urandom(16).hex()
            excel_sessions[session_id] = ExcelSession(
                params['filename'],
                params['saveLocation']
            )
            
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'session_id': session_id}).encode())

        except Exception as e:
            self.send_error(500, str(e))

    def handle_browse_directory(self):
        """Handles request to open a directory selection dialog."""
        try:
            # We need a Tk root window, but we don't want to show it
            root = tk.Tk()
            root.withdraw()  # Hide the main window
            root.attributes('-topmost', True) # Bring the dialog to the front

            # Open the directory selection dialog
            selected_path = filedialog.askdirectory(
                title="Select Save Location",
                initialdir=os.path.expanduser("~") # Start in user's home directory
            )
            root.destroy() # Close the Tkinter root window

            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            if selected_path:
                self.wfile.write(json.dumps({'path': selected_path}).encode())
            else:
                # User cancelled the dialog
                self.wfile.write(json.dumps({'path': None}).encode())

        except Exception as e:
            self.send_response(500)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e)}).encode())

    def get_upload_history(self):
        """Scans the UPLOAD_DIR for job folders and extracts history information."""
        history = []
        if not os.path.exists(UPLOAD_DIR):
            return history

        for job_id in os.listdir(UPLOAD_DIR):
            job_dir = os.path.join(UPLOAD_DIR, job_id)
            if os.path.isdir(job_dir):
                # Get directory creation time
                try:
                    created_time = os.path.getctime(job_dir)
                    created_date = datetime.fromtimestamp(created_time).strftime('%Y-%m-%d %H:%M:%S')
                except OSError:
                    created_date = "Unknown Date"

                job_type = "Unknown"
                input_files = []
                output_files = []

                # List all files in the job directory
                all_files_in_job_dir = [f for f in os.listdir(job_dir) if os.path.isfile(os.path.join(job_dir, f))]

                # --- Detect Job Type and Files ---
                is_video_job = any(f.lower().endswith(('.mp4', '.avi', '.mov', '.wmv', '.flv')) for f in all_files_in_job_dir) and \
                               any(f.lower().endswith(('.xlsx', '.xls')) for f in all_files_in_job_dir)
                is_word_job = any(f.lower().endswith(('.doc', '.docx')) for f in all_files_in_job_dir)

                if is_video_job:
                    job_type = "Video to PDF"
                    for f in all_files_in_job_dir:
                        if f.lower().endswith(('.mp4', '.avi', '.mov', '.wmv', '.flv', '.xlsx', '.xls')):
                            input_files.append(f)
                    # Look for output PDF in 'results' subdirectory
                    results_dir = os.path.join(job_dir, 'results')
                    if os.path.exists(results_dir):
                        for f in os.listdir(results_dir):
                            if f.lower().endswith('.pdf'):
                                output_files.append(f)
                elif is_word_job:
                    job_type = "Word to PDF"
                    for f in all_files_in_job_dir:
                        if f.lower().endswith(('.doc', '.docx')):
                            input_files.append(f)
                        elif f.lower().endswith('.pdf'): # PDF is directly in job_dir for Word conversion
                            output_files.append(f)
                # Add more job type detections here if needed (e.g., TTS)

                # Combine input and output files for display, prioritizing outputs if needed
                display_files = sorted(input_files) + sorted(output_files) # Simple concatenation for now

                if job_type != "Unknown" and display_files: # Only add jobs we could identify with files
                    history.append({
                        'id': job_id,
                        'date': created_date,
                        'type': job_type,
                        'files': display_files # Use the detected files
                    })

        # Sort by date, newest first
        history.sort(key=lambda x: x['date'], reverse=True)
        print(f"DEBUG: Generated history: {history}") # DEBUG - See what history is being generated
        return history

    def handle_get_tts_voices(self):
        """Handles GET request for available TTS voices."""
        if not TTS_ENABLED:
            self.send_error(501, "TTS functionality is not available.")
            return
        try:
            voices = get_available_voices() # Assuming this returns a list of dicts like [{'id': 'voice1', 'name': 'Voice One'}]
            self.send_response(200)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps(voices).encode())
        except Exception as e:
            self.send_error(500, f"Error getting TTS voices: {str(e)}")

    def handle_tts_conversion(self):
        """Handles POST request for TTS conversion."""
        if not TTS_ENABLED:
            self.send_error(501, "TTS functionality is not available.")
            return
        try:
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)
            data = json.loads(post_data)

            text = data.get('text')
            voice_id = data.get('voice_id')

            if not text or not voice_id:
                self.send_error(400, "Missing 'text' or 'voice_id' in request body")
                return

            # Assuming generate_tts returns the path to the generated MP3 file
            output_path = generate_tts(text, voice_id)

            if not output_path or not os.path.exists(output_path):
                 raise ValueError("TTS generation failed or did not produce an output file.")

            with open(output_path, 'rb') as f:
                audio_data = f.read()

            # Clean up the temporary audio file if necessary (depends on tts4.py implementation)
            # If generate_tts creates a temp file, uncomment the line below:
            # os.unlink(output_path)

            self.send_response(200)
            self.send_header('Content-type', 'audio/mpeg') # MP3 MIME type
            self.send_header('Content-Length', str(len(audio_data)))
            self.end_headers()
            self.wfile.write(audio_data)

        except Exception as e:
             error_message = f"Error during TTS conversion: {str(e)}"
             print(f"TTS Error: {error_message}") # Log the error server-side
             self.send_response(500)
             self.send_header('Content-type', 'application/json')
             self.end_headers()
             self.wfile.write(json.dumps({'error': error_message}).encode())

async def handle_websocket(websocket, path):
    try:
        # First message should be session ID
        session_id = await websocket.recv()
        session = excel_sessions.get(session_id)
        
        if not session:
            await websocket.close()
            return
        
        # Store websocket connection in session
        session.websocket = websocket
        session.monitoring = True
        
        # Start clipboard monitoring with delay
        def on_clipboard_change():
            if session.monitoring:
                # Wait a brief moment for the clipboard to update
                time.sleep(0.1)
                text = pyperclip.paste()
                asyncio.run(session.add_text(text))
        
        # Set up hotkey with delay
        keyboard.add_hotkey('ctrl+c', on_clipboard_change)
        
        try:
            while True:
                msg = await websocket.recv()
                data = json.loads(msg)
                
                if data['type'] == 'done':
                    session.monitoring = False
                    save_path = session.save()
                    await websocket.send(json.dumps({
                        'type': 'saved',
                        'path': save_path
                    }))
                    break
        
        finally:
            keyboard.remove_hotkey('ctrl+c')
            session.monitoring = False
            del excel_sessions[session_id]
    
    except websockets.exceptions.ConnectionClosed:
        pass

def run_websocket_server():
    asyncio.set_event_loop(asyncio.new_event_loop())
    start_server = websockets.serve(handle_websocket, "localhost", WS_PORT)
    asyncio.get_event_loop().run_until_complete(start_server)
    asyncio.get_event_loop().run_forever()

def run_http_server():
    with socketserver.TCPServer(("", HTTP_PORT), MultiToolHandler) as httpd:
        print(f"HTTP server running at http://localhost:{HTTP_PORT}")
        try:
            httpd.serve_forever()
        except KeyboardInterrupt:
            print("Server stopped by user")
            httpd.server_close()

if __name__ == "__main__":
    # Add MIME types
    if not mimetypes.guess_type('file.js')[0]:
        mimetypes.add_type('application/javascript', '.js')
    if not mimetypes.guess_type('file.css')[0]:
        mimetypes.add_type('text/css', '.css')
    
    # Start WebSocket server in a separate thread
    ws_thread = threading.Thread(target=run_websocket_server)
    ws_thread.daemon = True
    ws_thread.start()
    
    # Start HTTP server in main thread
    run_http_server()